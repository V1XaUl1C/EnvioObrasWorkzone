import streamlit as st
import pandas as pd
import os
import io
import math
import html
import re
import zipfile
import shutil
import stat
import tempfile
from datetime import datetime, timedelta

# Librerías para la Pestaña 4 (Versionador)
from docx import Document
import fitz  # PyMuPDF
import openpyxl # Reemplaza a win32com para hacer el código compatible con la Nube (Cloud)

# Configuración de página
st.set_page_config(layout="wide")

# ---------------- Funciones Auxiliares para Limpiar Interfaz ----------------
def recargar_pagina():
    if hasattr(st, "rerun"):
        st.rerun()
    else:
        st.experimental_rerun()

# ---------------- Archivos y Rutas ----------------
logo_path1 = os.path.join("data", "Logo_Pluz.png")
logo_path2 = os.path.join("data", "YoDoyEsePluz.png")
logo_path3 = os.path.join("data", "Portada.png")
base_datos_path = os.path.join("data", "Valorizaciones_Workzone.xlsx")
plantilla_base_path = os.path.join("data", "Plantilla_Base.xlsx") # TU PLANTILLA ORIGINAL

# ---------------- Logos y título ----------------
col1, col2, col3 = st.columns([1, 6, 1], vertical_alignment="center")
with col1:
    st.image(logo_path2, width=140)
with col2:
    st.markdown(
        """
        <h1 style="color:#2F56A6; margin:0; text-align:center;">
        Plantillas para valorizaciones - Proyectos de Distribución
        </h1>
        """,
        unsafe_allow_html=True
    )
with col3:
    st.image(logo_path1, width=140)

st.markdown("<hr style='border:2px solid #2F56A6;'>", unsafe_allow_html=True)

# ---------------- Funciones Base ----------------
def is_numeric(val):
    try:
        float(str(val).replace(',', '.'))
        return True
    except ValueError:
        return False

def get_val(df, col_search, val_search, col_return):
    if df is None or col_search not in df.columns or col_return not in df.columns:
        return ""
    match = df[df[col_search].astype(str).str.strip() == str(val_search).strip()]
    if not match.empty:
        resultado = match.iloc[0][col_return]
        if isinstance(resultado, pd.Series):
            return str(resultado.iloc[0]).strip()
        return str(resultado).strip()
    return ""

encabezados_salida = [
    "POSICION_NUMERO", "GRUPO_PRESUPUESTO", "POSICION_TIPO", "OPERACION_NUMERO", 
    "CONTRATO_ID", "CONTRATO_POSICION", "MATERIAL", "MATERIAL_UNIDAD_MEDIDA", 
    "MATERIAL_COSTO_UNITARIO", "MATERIAL_CANTIDAD", "MATERIAL_COSTO_TOTAL", 
    "FECHA_RESERVA", "EMPLAZAMIENTO", "PROY_INSPECTOR_USUARIO", "TENSION_NIVEL", 
    "PROY_TIPOLOGIA_ID", "SECCION", "PROY_PARTE_ID", "CODIGO_INTERNO", 
    "PROY_ZONA_ID", "PRESUPUESTO_FINALIDAD", "ELEMENTO_PEP", "ODM_PRESUPUESTO", 
    "ODM_PRESPTO_DESCRIPCION", "TAM"
]

lista_cx = ["CXEABT", "CXESBT", "CXGCAMT", "CXGCSBT", "CXGCSMT", "CXICBT", "CXITBT", "CXSMT", "CXABT", "CXAMT", "CXSBT"]

# =========================================================
# MOTOR DE PEGADO CRUDO (RAW PASTE) - Pestañas 1, 2, 3
# =========================================================
def generar_excel(dataframe):
    output = io.BytesIO()
    
    if os.path.exists(plantilla_base_path):
        wb = openpyxl.load_workbook(plantilla_base_path)
        ws = wb["PPTO."] if "PPTO." in wb.sheetnames else wb.active
        
        if ws.max_row > 1:
            ws.delete_rows(2, ws.max_row)
            
        for r_idx, row_data in enumerate(dataframe.to_dict('records'), start=2):
            for c_idx, header in enumerate(encabezados_salida, start=1):
                valor = row_data.get(header, "")
                ws.cell(row=r_idx, column=c_idx, value=valor)
                
        wb.save(output)
    else:
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            dataframe.to_excel(writer, index=False, sheet_name='PPTO.')
            
    return output.getvalue()

# =========================================================
# TABLAS HTML ELEGANTES
# =========================================================
def mostrar_tabla_bonita(df, height="350px"):
    html_code = df.to_html(index=False)
    html_code = html_code.replace('<table border="1" class="dataframe">', '<table style="width:100%; border-collapse: collapse; font-family: sans-serif; font-size: 13px;">')
    html_code = html_code.replace('<th>', '<th style="background-color: #2F56A6; color: white; padding: 10px; border: 1px solid #ddd; text-align: center; position: sticky; top: 0; z-index: 1;">')
    html_code = html_code.replace('<td>', '<td style="padding: 8px; border: 1px solid #ddd; text-align: center;">')
    
    st.markdown(
        f'<div style="max-height: {height}; overflow-y: auto; overflow-x: auto; border: 1px solid #ddd; border-radius: 5px; margin-bottom: 20px;">{html_code}</div>', 
        unsafe_allow_html=True
    )

@st.cache_data
def cargar_bases_datos(ruta):
    def arreglar_encabezados(df, columna_clave):
        df.columns = [str(c).strip() for c in df.columns]
        if columna_clave in df.columns: return df
        for i, row in df.iterrows():
            valores_fila = [str(x).strip() for x in row.values]
            if columna_clave in valores_fila:
                df.columns = valores_fila
                return df.iloc[i+1:].reset_index(drop=True)
        return df

    try:
        df_guia = arreglar_encabezados(pd.read_excel(ruta, sheet_name="Guia_SED"), 'SED')
        df_capex = arreglar_encabezados(pd.read_excel(ruta, sheet_name="Capex"), 'Cod. Int.')
        df_datos = arreglar_encabezados(pd.read_excel(ruta, sheet_name="Datos"), 'DISTRITO')
        df_mat = arreglar_encabezados(pd.read_excel(ruta, sheet_name="Materiales"), 'Material SAP')
        df_serv = arreglar_encabezados(pd.read_excel(ruta, sheet_name="Servicios"), 'Ser-Mat ASIS')
        df_peps = pd.read_excel(ruta, sheet_name="PEPs x FB Técnica Dx", header=None)
        return df_guia, df_capex, df_datos, df_mat, df_serv, df_peps
    except Exception as e:
        st.error(f"Error crítico al leer Excel BD: {e}")
        return None, None, None, None, None, None

df_guia, df_capex, df_datos, df_mat, df_serv, df_peps = cargar_bases_datos(base_datos_path)

# =========================================================
# FUNCIONES DE PROCESAMIENTO (PESTAÑA 4) - MEJORADAS
# =========================================================
def procesar_word(ruta, viejo, nuevo):
    doc = Document(ruta)
    for p in doc.paragraphs:
        if viejo in p.text:
            p.text = p.text.replace(viejo, nuevo)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if viejo in celda.text:
                    celda.text = celda.text.replace(viejo, nuevo)
    doc.save(ruta)

def procesar_excel_automatico(ruta_absoluta, viejo, nuevo):
    try:
        wb = openpyxl.load_workbook(ruta_absoluta)
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        # Si es texto normal
                        if isinstance(cell.value, str):
                            if viejo in cell.value:
                                cell.value = cell.value.replace(viejo, nuevo)
                        # Si es un número (y el código viejo era numérico)
                        elif isinstance(cell.value, (int, float)):
                            val_str = str(int(cell.value)) if isinstance(cell.value, float) and cell.value.is_integer() else str(cell.value)
                            if viejo in val_str:
                                nuevo_val = val_str.replace(viejo, nuevo)
                                try:
                                    cell.value = int(nuevo_val)
                                except ValueError:
                                    cell.value = nuevo_val
        wb.save(ruta_absoluta)
        return True
    except Exception as e:
        print(f"Error detallado con Excel: {e}")
        return False

def procesar_pdf(ruta, viejo, nuevo):
    doc = fitz.open(ruta)
    for pagina in doc:
        areas_texto = pagina.search_for(viejo)
        spans_formato = []
        for rect in areas_texto:
            info_texto = pagina.get_text("dict", clip=rect)
            tamano_fuente = 12
            color_fuente = (0,0,0)
            try:
                for bloque in info_texto.get("blocks", []):
                    for linea in bloque.get("lines", []):
                        for span in linea.get("spans", []):
                            if viejo in span.get("text",""):
                                tamano_fuente = span["size"]
                                color_int = span["color"]
                                r = ((color_int>>16)&255)/255
                                g = ((color_int>>8)&255)/255
                                b = (color_int&255)/255
                                color_fuente = (r,g,b)
                                raise StopIteration
            except StopIteration:
                pass
            spans_formato.append((rect, tamano_fuente, color_fuente))
            pagina.add_redact_annot(rect, fill=(1,1,1))
        pagina.apply_redactions()

        for rect, tamano_fuente, color_fuente in spans_formato:
            punto_insercion = fitz.Point(rect.x0, rect.y1 - (rect.height*0.15))
            pagina.insert_text(punto_insercion, nuevo, fontsize=tamano_fuente, color=color_fuente, fontname="helv")

    ruta_temp = ruta+".tmp"
    doc.save(ruta_temp, incremental=False, encryption=0)
    doc.close()
    os.replace(ruta_temp, ruta)

# ---------------- Interfaz con Pestañas ----------------
st.markdown("<h3 style='color:#2F56A6;'>Selecciona el módulo a utilizar:</h3>", unsafe_allow_html=True)

tab1, tab2, tab3, tab4 = st.tabs([
    "🛠️ Plantilla ODM ING", 
    "🔄 Convertir archivo VAS", 
    "📝 Convertir plantilla OT",
    "📁 Versionar documentos"
])

# --- SECCIÓN 1: Plantilla ODM ING ---
with tab1:
    st.subheader("Generación de Plantilla ODM ING")

    if df_guia is not None:
        col_input1, col_input2, col_input3 = st.columns(3)
        with col_input1:
            lista_seds = df_guia['SED'].dropna().unique().tolist() if 'SED' in df_guia.columns else []
            sed_seleccionada = st.selectbox("1. Selecciona SED", options=[""] + lista_seds, key="t1_sed")
        with col_input2:
            lista_codigos = df_capex['Cod. Int.'].dropna().astype(str).unique().tolist() if 'Cod. Int.' in df_capex.columns else []
            codigo_seleccionado = st.selectbox("2. Selecciona Código Interno", options=[""] + lista_codigos, key="t1_cod")
        with col_input3:
            tension_seleccionada = st.selectbox("3. Selecciona Nivel de Tensión", options=["", "BT", "MT"], key="t1_ten")
        
        # --- MODIFICACIÓN: Nuevos inputs Área y Contrato ---
        col_input4, col_input5, _ = st.columns([1, 1, 1])
        with col_input4:
            area_seleccionada = st.selectbox("4. Área", options=["Proyectos BT", "Proyectos MT/BT"], key="t1_area")
        with col_input5:
            contrato_seleccionado = st.selectbox("5. Contrato", options=["Automático", "Applus Colonial", "Applus Panamericana", "Satel"], key="t1_con")
        st.markdown("---")
        
        col_btn1, col_btn2 = st.columns([2, 8])
        
        if col_btn2.button("🧹 Limpiar Campos", key="btn_limpiar_tab1"):
            for k in ['t1_sed', 't1_cod', 't1_ten', 't1_area', 't1_con']:
                st.session_state.pop(k, None)
            recargar_pagina()

        if col_btn1.button("🚀 Generar Plantilla Excel"):
            if not sed_seleccionada or not codigo_seleccionado or not tension_seleccionada:
                st.warning("⚠️ Completa los campos.")
            else:
                fila_sed = df_guia[df_guia['SED'] == sed_seleccionada].iloc[0]
                alim = str(fila_sed.get('ALIM', '')).strip()
                org_mantenimiento = str(fila_sed.get('ORGANIZACIÓN MANTENIMIENTO', '')).strip().upper()
                distrito = str(fila_sed.get('DISTRITO', '')).strip().upper()
                
                # --- LÓGICA DE SELECCIÓN DE CONTRATO ---
                if contrato_seleccionado == "Automático":
                    if area_seleccionada == "Proyectos BT":
                        if distrito in ["LOS OLIVOS", "COMAS"]:
                            contrato_id = "5200000072"
                        elif "COLONIAL" in org_mantenimiento:
                            contrato_id = "5200000075"
                        elif "PANAMERICANA" in org_mantenimiento:
                            contrato_id = "5200000102"
                        else:
                            contrato_id = "" 
                    elif area_seleccionada == "Proyectos MT/BT":
                        contrato_id = "5200000072"
                    else:
                        contrato_id = ""
                elif contrato_seleccionado == "Applus Colonial":
                    contrato_id = "5200000075"
                elif contrato_seleccionado == "Applus Panamericana":
                    contrato_id = "5200000102"
                elif contrato_seleccionado == "Satel":
                    contrato_id = "5200000072"
                else:
                    contrato_id = ""

                zona_id = "A" if "COLONIAL" in org_mantenimiento else ("B" if "PANAMERICANA" in org_mantenimiento else "")
                fecha_reserva = datetime.now().strftime("%d.%m.%Y")
                emplazamiento = f"M-{alim}"
                
                fb_finalidad_ing = get_val(df_capex, 'Cod. Int.', codigo_seleccionado, 'FB')
                
                fila_base = {
                    "GRUPO_PRESUPUESTO": "ING", "OPERACION_NUMERO": 10, "CONTRATO_ID": contrato_id, "CONTRATO_POSICION": "",
                    "MATERIAL_UNIDAD_MEDIDA": "", "MATERIAL_COSTO_UNITARIO": "", "MATERIAL_CANTIDAD": 1, "MATERIAL_COSTO_TOTAL": "",
                    "FECHA_RESERVA": fecha_reserva, "EMPLAZAMIENTO": emplazamiento, "PROY_INSPECTOR_USUARIO": "", "TENSION_NIVEL": tension_seleccionada,
                    "PROY_TIPOLOGIA_ID": "N", "SECCION": 154, "PROY_PARTE_ID": 1545, "CODIGO_INTERNO": codigo_seleccionado, "PROY_ZONA_ID": zona_id,
                    "PRESUPUESTO_FINALIDAD": fb_finalidad_ing, "ELEMENTO_PEP": "", "ODM_PRESUPUESTO": "", "ODM_PRESPTO_DESCRIPCION": "Ingenieria de Proyectos", "TAM": ""
                }
                
                fila1 = fila_base.copy()
                fila1.update({"POSICION_NUMERO": 1, "POSICION_TIPO": "N", "MATERIAL": "2100111"})
                
                fila2 = fila_base.copy()
                fila2.update({"POSICION_NUMERO": 2, "POSICION_TIPO": "L", "MATERIAL": "S2100111"})
                
                df_resultado = pd.DataFrame([fila1, fila2], columns=encabezados_salida)
                st.write("### Vista previa:")
                mostrar_tabla_bonita(df_resultado, height="150px")
                st.download_button("📥 Descargar Plantilla", data=generar_excel(df_resultado), file_name=f"Plantilla_ODM_ING_{codigo_seleccionado}.xlsx")

# --- SECCIÓN 2: Convertir archivo VAS ---
with tab2:
    st.subheader("Conversión de archivo VAS a Plantilla Final")

    # Clave dinámica para resetear el uploader
    if 'vas_uploader_key' not in st.session_state: 
        st.session_state['vas_uploader_key'] = 0

    col_vas1, col_vas2, col_vas3, col_vas4, col_vas5 = st.columns(5)
    with col_vas1:
        archivo_vas = st.file_uploader("Sube el archivo VAS (.xlsx)", type=['xlsx'], key=f"file_vas_{st.session_state['vas_uploader_key']}")
        
        if archivo_vas is None:
            if 'datos_vas' in st.session_state: del st.session_state['datos_vas']
            if 'params_vas' in st.session_state: del st.session_state['params_vas']

    with col_vas2:
        seds_disponibles = df_guia['SED'].dropna().unique().tolist() if df_guia is not None and 'SED' in df_guia.columns else []
        sed_vas = st.selectbox("1. SED", options=[""] + seds_disponibles, key="sed_vas2")
    with col_vas3:
        codigos_disponibles = df_capex['Cod. Int.'].dropna().astype(str).unique().tolist() if df_capex is not None and 'Cod. Int.' in df_capex.columns else []
        cod_pry_vas = st.selectbox("2. Código PRY", options=[""] + codigos_disponibles, key="cod_pry_vas2")
    with col_vas4:
        plazo_vas = st.selectbox("3. Plazo Legal", options=["", "21", "56", "360"], key="plazo_vas_key")
    with col_vas5:
        tipologia_vas = st.selectbox("4. Tipología ID", options=["", "N", "R"], key="tipo_vas_key")

    st.markdown("---")
    
    col_b1, col_b2 = st.columns([2, 8])
    btn_procesar_vas = col_b1.button("⚙️ Procesar VAS y Armar Plantilla")

    # --- BOTÓN LIMPIAR VAS ---
    if col_b2.button("🧹 Limpiar Todo", key="btn_limpiar_tab2"):
        st.session_state['vas_uploader_key'] += 1
        for k in ['sed_vas2', 'cod_pry_vas2', 'plazo_vas_key', 'tipo_vas_key', 'datos_vas', 'params_vas']:
            st.session_state.pop(k, None)
        recargar_pagina()

    if btn_procesar_vas:
        if not archivo_vas or not sed_vas or not cod_pry_vas or not plazo_vas or not tipologia_vas:
            st.warning("⚠️ Por favor, sube el archivo y completa todos los parámetros.")
        else:
            try:
                df_vas = pd.read_excel(archivo_vas, header=None, dtype=str)
                filas = df_vas.fillna("").values.tolist()

                cit_to_agp = {} 
                bloque_materiales, bloque_servicios, bloque_espejos = [], [], []
                idx_agp, idx_cit = -1, -1
                idx_tipo, idx_contador, idx_mat, idx_cant = -1, -1, -1, -1
                
                distrito = get_val(df_guia, 'SED', sed_vas, 'DISTRITO')
                
                for fila in filas:
                    cols = [str(c).strip() for c in fila]
                    if not any(cols): continue
                    
                    if "Agrup.Prot.AGP" in cols:
                        idx_agp = cols.index("Agrup.Prot.AGP")
                        idx_cit = cols.index("CIT") if "CIT" in cols else min(3, len(cols)-1)
                        continue

                    if "Mat./Prest." in cols and "Tipo" in cols:
                        idx_tipo = cols.index("Tipo")
                        idx_contador = cols.index("Contador") if "Contador" in cols else -1
                        idx_mat = cols.index("Mat./Prest.")
                        idx_cant = cols.index("Cantidad") if "Cantidad" in cols else -1
                        continue
                    
                    if idx_agp != -1 and idx_cit != -1 and len(cols) > max(idx_agp, idx_cit):
                        val_agp, val_cit = cols[idx_agp], cols[idx_cit]
                        if val_agp and val_cit and val_agp != "Agrup.Prot.AGP" and "IMPORTE" not in val_agp.upper():
                            if not (idx_tipo != -1 and len(cols) > idx_tipo and cols[idx_tipo] in ["Materiales", "Servicios"]):
                                cit_to_agp[val_cit] = val_agp
                                
                    if idx_tipo != -1 and idx_mat != -1 and idx_cant != -1 and idx_contador != -1:
                        if len(cols) > max(idx_tipo, idx_contador, idx_mat, idx_cant):
                            tipo = cols[idx_tipo]
                            if tipo in ["Materiales", "Servicios"]:
                                contador = cols[idx_contador]
                                matricula = cols[idx_mat]
                                cantidad_str = cols[idx_cant].replace(',', '.')
                                
                                cit_val = contador.split('-')[-1] if '-' in contador else contador
                                agrupador = cit_to_agp.get(cit_val, "Desconocido")
                                
                                if matricula.endswith('.0'): matricula = matricula[:-2]
                                if tipo == "Servicios" and len(matricula) >= 3: matricula = "TLA" + matricula[3:]
                                    
                                cantidad_final = math.ceil(float(cantidad_str)) if (tipo == "Servicios" and is_numeric(cantidad_str)) else (float(cantidad_str) if is_numeric(cantidad_str) else cantidad_str)
                                    
                                row_data = {
                                    "Agrup.Prot.AGP": agrupador, "Tipo": tipo,
                                    "Mat./Prest.": matricula, "Cantidad": cantidad_final
                                }
                                
                                if tipo == "Materiales": bloque_materiales.append(row_data)
                                elif tipo == "Servicios":
                                    bloque_servicios.append(row_data)
                                    row_espejo = row_data.copy()
                                    row_espejo["Tipo"] = "Espejo"
                                    bloque_espejos.append(row_espejo)

                st.session_state['datos_vas'] = bloque_materiales + bloque_servicios + bloque_espejos
                st.session_state['params_vas'] = {
                    "sed": sed_vas, "pry": cod_pry_vas, "plazo": plazo_vas, "tipologia": tipologia_vas, "distrito": distrito
                }
            except Exception as e:
                st.error(f"❌ Error leyendo el archivo: {e}")

    if 'datos_vas' in st.session_state:
        datos_vas = st.session_state['datos_vas']
        params = st.session_state['params_vas']
        fb_finalidad = get_val(df_capex, 'Cod. Int.', params['pry'], 'FB')
        
        agrupadores_unicos = {row['Agrup.Prot.AGP'] for row in datos_vas}
        agrupadores_con_vacios = []
        peps_validos = []
        if df_peps is not None:
            mask = df_peps[5].astype(str).str.strip() == str(fb_finalidad).strip()
            peps_validos = df_peps[mask][1].astype(str).str.strip().dropna().unique().tolist()
            peps_validos = [p for p in peps_validos if p and p.lower() != 'nan']
            
            peps_h3 = df_peps[df_peps[5].astype(str).str.strip() == "H3"][1].astype(str).str.strip().dropna().unique().tolist()
            peps_especiales = [p for p in peps_h3 if p in ["RAAP", "RSAP"]]
            for especial in peps_especiales:
                if especial not in peps_validos:
                    peps_validos.append(especial)

        for agp in agrupadores_unicos:
            if agp not in lista_cx:
                fb_check = "H3" if agp in ["RAAP", "RSAP"] else fb_finalidad
                filtro_pep = df_peps[(df_peps[1].astype(str).str.strip() == str(agp)) & (df_peps[5].astype(str).str.strip() == str(fb_check))]
                if filtro_pep.empty:
                    agrupadores_con_vacios.append(agp)

        mapeo_usuario = {}
        if agrupadores_con_vacios:
            st.error("⚠️ Atención: Identificamos Agrupadores que dejarán Sección y Parte vacíos. Por favor, selecciona el correcto:")
            cols_map = st.columns(3)
            for i, agp_vacio in enumerate(agrupadores_con_vacios):
                opciones_combinadas = lista_cx + peps_validos
                opciones = ["(Dejar vacío)"] + opciones_combinadas
                
                sugerencia_idx = 0
                for j, opcion_valida in enumerate(opciones_combinadas):
                    if agp_vacio.startswith(opcion_valida) or opcion_valida.startswith(agp_vacio[:4]):
                        sugerencia_idx = j + 1
                        break
                
                with cols_map[i % 3]:
                    seleccion = st.selectbox(f"Corregir '{agp_vacio}' por:", options=opciones, index=sugerencia_idx, key=f"corrige_vas_{agp_vacio}")
                    if seleccion != "(Dejar vacío)": mapeo_usuario[agp_vacio] = seleccion

        st.markdown("### Plantilla Final (VAS)")

        datos_finales = []
        alim = get_val(df_guia, 'SED', params['sed'], 'ALIM')
        org_mant = str(get_val(df_guia, 'SED', params['sed'], 'ORGANIZACIÓN MANTENIMIENTO')).upper()
        contrato_final = get_val(df_datos, 'DISTRITO', params['distrito'], 'CONTRATO')
        zona_final = "A" if "COLONIAL" in org_mant else ("B" if "PANAMERICANA" in org_mant else "")
        emplazamiento_final = f"M-{alim}" if alim else ""
        dias_sumar = 3 if params['plazo'] == "21" else (5 if params['plazo'] == "56" else (15 if params['plazo'] == "360" else 0))
        fecha_reserva_final = (datetime.now() + timedelta(days=dias_sumar)).strftime("%d.%m.%Y")
        
        for i, row in enumerate(datos_vas):
            agrupador_original = row["Agrup.Prot.AGP"]
            agrupador_busqueda = mapeo_usuario.get(agrupador_original, agrupador_original)
            tipo, matricula, cantidad = row["Tipo"], row["Mat./Prest."], row["Cantidad"]
            
            mat_final = ""
            if tipo == "Materiales": mat_final = get_val(df_mat, 'Material SAP', matricula, 'Matricula EON TOBE')
            elif tipo == "Servicios": mat_final = get_val(df_serv, 'Ser-Mat ASIS', matricula, 'Ser-Mat TOBE')
            elif tipo == "Espejo":
                hallazgo = get_val(df_serv, 'Ser-Mat ASIS', matricula, 'Ser-Mat TOBE')
                mat_final = f"S{hallazgo}" if hallazgo else ""
                
            posicion_tipo = "N" if str(mat_final).startswith("2") else "L"
            seccion, parte, descripcion = "", "", ""
            if agrupador_busqueda in lista_cx:
                tension = get_val(df_datos, 'GRUPO_PRESUPUESTO', agrupador_busqueda, 'TENSION_NIVEL')
                seccion = get_val(df_datos, 'GRUPO_PRESUPUESTO', agrupador_busqueda, 'SECCION')
                parte = get_val(df_datos, 'GRUPO_PRESUPUESTO', agrupador_busqueda, 'PROY_PARTE_ID')
                descripcion = get_val(df_datos, 'GRUPO_PRESUPUESTO', agrupador_busqueda, 'DESCRIPCION DEL PEP')
            else:
                tension = "MT" if str(agrupador_busqueda).startswith(("SE", "RADP", "RSDP")) else "BT"
                fb_temp = "H3" if agrupador_busqueda in ["RAAP", "RSAP"] else fb_finalidad
                filtro_pep = df_peps[(df_peps[1].astype(str).str.strip() == str(agrupador_busqueda)) & (df_peps[5].astype(str).str.strip() == str(fb_temp))]
                if not filtro_pep.empty:
                    seccion, parte = filtro_pep.iloc[0][18], filtro_pep.iloc[0][20]   
                match_desc = df_peps[df_peps[1].astype(str).str.strip() == str(agrupador_busqueda)]
                if not match_desc.empty: descripcion = match_desc.iloc[0][21]

            # --- LÓGICA DE DEFINICIÓN DE CÓDIGO (RAAP/RSAP y DS11) ---
            pry_str = str(params['pry']).strip().upper()
            if agrupador_busqueda in ["RAAP", "RSAP"]:
                fb_row = "H3"
                codigo_row = "AP100" if pry_str != "DD001" else "AP053"
            else:
                fb_row = fb_finalidad
                if pry_str.startswith("DS11"):
                    desc_upper = str(descripcion).upper()
                    if str(agrupador_busqueda).upper().startswith("SE"):
                        codigo_row = "DS113"
                    elif "MT" in desc_upper:
                        codigo_row = "DS112"
                    elif "BT" in desc_upper:
                        codigo_row = "DS111"
                    else:
                        codigo_row = pry_str
                else:
                    codigo_row = pry_str

            fila_excel = {
                "POSICION_NUMERO": i + 1, 
                "GRUPO_PRESUPUESTO": agrupador_busqueda,
                "POSICION_TIPO": posicion_tipo,
                "OPERACION_NUMERO": 10, "CONTRATO_ID": contrato_final, "CONTRATO_POSICION": "", "MATERIAL": mat_final,
                "MATERIAL_UNIDAD_MEDIDA": "", "MATERIAL_COSTO_UNITARIO": "", "MATERIAL_CANTIDAD": cantidad, "MATERIAL_COSTO_TOTAL": "",
                "FECHA_RESERVA": fecha_reserva_final, "EMPLAZAMIENTO": emplazamiento_final, "PROY_INSPECTOR_USUARIO": "", "TENSION_NIVEL": tension,
                "PROY_TIPOLOGIA_ID": params['tipologia'], "SECCION": seccion, "PROY_PARTE_ID": parte, 
                "CODIGO_INTERNO": codigo_row,
                "PROY_ZONA_ID": zona_final, 
                "PRESUPUESTO_FINALIDAD": fb_row,
                "ELEMENTO_PEP": "", "ODM_PRESUPUESTO": "", "ODM_PRESPTO_DESCRIPCION": descripcion, "TAM": ""
            }
            datos_finales.append(fila_excel)

        df_final = pd.DataFrame(datos_finales)
        mostrar_tabla_bonita(df_final)
        
        st.download_button(
            label="📥 Descargar Plantilla Completada (.xlsx)", data=generar_excel(df_final),
            file_name=f"Plantilla_VAS_{params['pry']}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

# --- SECCIÓN 3: Convertir plantilla OT ---
with tab3:
    st.subheader("Conversión de plantilla OT a Plantilla Final")
    st.info("Sube tu archivo Pre-Matriz (Excel) con las columnas: Agrup.Prot.AGP, Tipo, Mat./LrAst., Cantidad.")

    if 'ot_uploader_key' not in st.session_state: 
        st.session_state['ot_uploader_key'] = 0

    col_ot1, col_ot2, col_ot3, col_ot4, col_ot5 = st.columns(5)
    with col_ot1:
        archivo_ot = st.file_uploader("Sube el archivo OT (.xlsx)", type=['xlsx'], key=f"file_ot_{st.session_state['ot_uploader_key']}")

        if archivo_ot is None:
            if 'datos_ot' in st.session_state: del st.session_state['datos_ot']
            if 'params_ot' in st.session_state: del st.session_state['params_ot']

    with col_ot2:
        seds_disponibles = df_guia['SED'].dropna().unique().tolist() if df_guia is not None and 'SED' in df_guia.columns else []
        sed_ot = st.selectbox("1. SED", options=[""] + seds_disponibles, key="sed_ot")
    with col_ot3:
        codigos_disponibles = df_capex['Cod. Int.'].dropna().astype(str).unique().tolist() if df_capex is not None and 'Cod. Int.' in df_capex.columns else []
        cod_pry_ot = st.selectbox("2. Código PRY", options=[""] + codigos_disponibles, key="cod_pry_ot")
    with col_ot4:
        plazo_ot = st.selectbox("3. Plazo Legal", options=["", "21", "56", "360"], key="plazo_ot_key")
    with col_ot5:
        tipologia_ot = st.selectbox("4. Tipología ID", options=["", "N", "R"], key="tipo_ot_key")

    st.markdown("---")

    col_b1, col_b2 = st.columns([2, 8])
    btn_procesar_ot = col_b1.button("⚙️ Procesar OT y Armar Plantilla")

    # --- BOTÓN LIMPIAR OT ---
    if col_b2.button("🧹 Limpiar Todo", key="btn_limpiar_tab3"):
        st.session_state['ot_uploader_key'] += 1
        for k in ['sed_ot', 'cod_pry_ot', 'plazo_ot_key', 'tipo_ot_key', 'datos_ot', 'params_ot']:
            st.session_state.pop(k, None)
        recargar_pagina()

    if btn_procesar_ot:
        if not archivo_ot or not sed_ot or not cod_pry_ot or not plazo_ot or not tipologia_ot:
            st.warning("⚠️ Por favor, sube el archivo OT y completa todos los parámetros.")
        else:
            try:
                df_ot_uploaded = pd.read_excel(archivo_ot)
                cols_dict = {str(c).strip().lower(): str(c) for c in df_ot_uploaded.columns}
                
                col_agp = cols_dict.get("agrup.prot.agp")
                col_tipo = cols_dict.get("tipo")
                col_mat = cols_dict.get("mat./lrast.") or cols_dict.get("mat./prest.") 
                col_cant = cols_dict.get("cantidad")

                if not (col_agp and col_tipo and col_mat and col_cant):
                    st.error(f"❌ El archivo no tiene las columnas correctas.")
                else:
                    distrito_ot = get_val(df_guia, 'SED', sed_ot, 'DISTRITO')
                    bloque_mat, bloque_serv, bloque_esp = [], [], []
                    tiene_espejos_previos = 'espejo' in df_ot_uploaded[col_tipo].astype(str).str.lower().values
                    
                    for _, row in df_ot_uploaded.iterrows():
                        agrupador = str(row[col_agp]).strip()
                        tipo = str(row[col_tipo]).strip()
                        matricula = str(row[col_mat]).strip()
                        cantidad_str = str(row[col_cant]).replace(',', '.').strip()
                        
                        if not agrupador or agrupador == "nan": continue
                            
                        if matricula.endswith('.0'): matricula = matricula[:-2]
                        
                        if tipo == "Servicios" and len(matricula) >= 3 and not matricula.startswith("TLA"): 
                            matricula = "TLA" + matricula[3:]
                            
                        cantidad_final = math.ceil(float(cantidad_str)) if (tipo == "Servicios" and is_numeric(cantidad_str)) else (float(cantidad_str) if is_numeric(cantidad_str) else cantidad_str)
                            
                        row_data = {
                            "Agrup.Prot.AGP": agrupador, "Tipo": tipo,
                            "Mat./Prest.": matricula, "Cantidad": cantidad_final
                        }
                        
                        if tipo == "Materiales": bloque_mat.append(row_data)
                        elif tipo == "Servicios": 
                            bloque_serv.append(row_data)
                            if not tiene_espejos_previos:
                                row_espejo = row_data.copy()
                                row_espejo["Tipo"] = "Espejo"
                                bloque_esp.append(row_espejo)
                        elif tipo == "Espejo":
                            bloque_esp.append(row_data)

                    st.session_state['datos_ot'] = bloque_mat + bloque_serv + bloque_esp
                    st.session_state['params_ot'] = {
                       "sed": sed_ot, "pry": cod_pry_ot, "plazo": plazo_ot, "tipologia": tipologia_ot, "distrito": distrito_ot
                    }
            except Exception as e:
                st.error(f"❌ Error leyendo el archivo OT: {e}")

    if 'datos_ot' in st.session_state:
        datos_ot = st.session_state['datos_ot']
        params = st.session_state['params_ot']
        fb_finalidad = get_val(df_capex, 'Cod. Int.', params['pry'], 'FB')
        
        agrupadores_unicos = {row['Agrup.Prot.AGP'] for row in datos_ot}
        agrupadores_con_vacios = []
        peps_validos = []
        if df_peps is not None:
            mask = df_peps[5].astype(str).str.strip() == str(fb_finalidad).strip()
            peps_validos = df_peps[mask][1].astype(str).str.strip().dropna().unique().tolist()
            peps_validos = [p for p in peps_validos if p and p.lower() != 'nan']
            
            peps_h3 = df_peps[df_peps[5].astype(str).str.strip() == "H3"][1].astype(str).str.strip().dropna().unique().tolist()
            peps_especiales = [p for p in peps_h3 if p in ["RAAP", "RSAP"]]
            for especial in peps_especiales:
                if especial not in peps_validos:
                    peps_validos.append(especial)

        for agp in agrupadores_unicos:
            if agp not in lista_cx:
                fb_check = "H3" if agp in ["RAAP", "RSAP"] else fb_finalidad
                filtro_pep = df_peps[(df_peps[1].astype(str).str.strip() == str(agp)) & (df_peps[5].astype(str).str.strip() == str(fb_check))]
                if filtro_pep.empty:
                    agrupadores_con_vacios.append(agp)

        mapeo_usuario_ot = {}
        if agrupadores_con_vacios:
            st.error("⚠️ Atención: Identificamos Agrupadores que dejarán Sección y Parte vacíos. Por favor, selecciona el correcto:")
            cols_map = st.columns(3)
            for i, agp_vacio in enumerate(agrupadores_con_vacios):
                opciones_combinadas_ot = lista_cx + peps_validos
                opciones = ["(Dejar vacío)"] + opciones_combinadas_ot
                
                sugerencia_idx = 0
                for j, opcion_valida in enumerate(opciones_combinadas_ot):
                    if agp_vacio.startswith(opcion_valida) or opcion_valida.startswith(agp_vacio[:4]):
                        sugerencia_idx = j + 1
                        break
                
                with cols_map[i % 3]:
                    seleccion = st.selectbox(f"Corregir '{agp_vacio}' por:", options=opciones, index=sugerencia_idx, key=f"corrige_ot_{agp_vacio}")
                    if seleccion != "(Dejar vacío)": mapeo_usuario_ot[agp_vacio] = seleccion

        st.markdown("### Plantilla Final (OT)")

        datos_finales_ot = []
        alim = get_val(df_guia, 'SED', params['sed'], 'ALIM')
        org_mant = str(get_val(df_guia, 'SED', params['sed'], 'ORGANIZACIÓN MANTENIMIENTO')).upper()
        contrato_final = get_val(df_datos, 'DISTRITO', params['distrito'], 'CONTRATO')
        zona_final = "A" if "COLONIAL" in org_mant else ("B" if "PANAMERICANA" in org_mant else "")
        emplazamiento_final = f"M-{alim}" if alim else ""
        dias_sumar = 3 if params['plazo'] == "21" else (5 if params['plazo'] == "56" else (15 if params['plazo'] == "360" else 0))
        fecha_reserva_final = (datetime.now() + timedelta(days=dias_sumar)).strftime("%d.%m.%Y")
        
        for i, row in enumerate(datos_ot):
            agrupador_original = row["Agrup.Prot.AGP"]
            agrupador_busqueda = mapeo_usuario_ot.get(agrupador_original, agrupador_original)
            tipo, matricula, cantidad = row["Tipo"], row["Mat./Prest."], row["Cantidad"]
            
            mat_final = ""
            if tipo == "Materiales": mat_final = get_val(df_mat, 'Material SAP', matricula, 'Matricula EON TOBE')
            elif tipo == "Servicios": mat_final = get_val(df_serv, 'Ser-Mat ASIS', matricula, 'Ser-Mat TOBE')
            elif tipo == "Espejo":
                hallazgo = get_val(df_serv, 'Ser-Mat ASIS', matricula, 'Ser-Mat TOBE')
                mat_final = f"S{hallazgo}" if hallazgo else ""
                
            posicion_tipo = "N" if str(mat_final).startswith("2") else "L"
            seccion, parte, descripcion = "", "", ""
            if agrupador_busqueda in lista_cx:
                tension = get_val(df_datos, 'GRUPO_PRESUPUESTO', agrupador_busqueda, 'TENSION_NIVEL')
                seccion = get_val(df_datos, 'GRUPO_PRESUPUESTO', agrupador_busqueda, 'SECCION')
                parte = get_val(df_datos, 'GRUPO_PRESUPUESTO', agrupador_busqueda, 'PROY_PARTE_ID')
                descripcion = get_val(df_datos, 'GRUPO_PRESUPUESTO', agrupador_busqueda, 'DESCRIPCION DEL PEP')
            else:
                tension = "MT" if str(agrupador_busqueda).startswith(("SE", "RADP", "RSDP")) else "BT"
                fb_temp = "H3" if agrupador_busqueda in ["RAAP", "RSAP"] else fb_finalidad
                filtro_pep = df_peps[(df_peps[1].astype(str).str.strip() == str(agrupador_busqueda)) & (df_peps[5].astype(str).str.strip() == str(fb_temp))]
                if not filtro_pep.empty:
                    seccion, parte = filtro_pep.iloc[0][18], filtro_pep.iloc[0][20]   
                match_desc = df_peps[df_peps[1].astype(str).str.strip() == str(agrupador_busqueda)]
                if not match_desc.empty: descripcion = match_desc.iloc[0][21]

            # --- LÓGICA DE DEFINICIÓN DE CÓDIGO (RAAP/RSAP y DS11) ---
            pry_str = str(params['pry']).strip().upper()
            if agrupador_busqueda in ["RAAP", "RSAP"]:
                fb_row = "H3"
                codigo_row = "AP100" if pry_str != "DD001" else "AP053"
            else:
                fb_row = fb_finalidad
                if pry_str.startswith("DS11"):
                    desc_upper = str(descripcion).upper()
                    if str(agrupador_busqueda).upper().startswith("SE"):
                        codigo_row = "DS113"
                    elif "MT" in desc_upper:
                        codigo_row = "DS112"
                    elif "BT" in desc_upper:
                        codigo_row = "DS111"
                    else:
                        codigo_row = pry_str
                else:
                    codigo_row = pry_str

            fila_excel = {
                "POSICION_NUMERO": i + 1, 
                "GRUPO_PRESUPUESTO": agrupador_busqueda,
                "POSICION_TIPO": posicion_tipo,
                "OPERACION_NUMERO": 10, "CONTRATO_ID": contrato_final, "CONTRATO_POSICION": "", "MATERIAL": mat_final,
                "MATERIAL_UNIDAD_MEDIDA": "", "MATERIAL_COSTO_UNITARIO": "", "MATERIAL_CANTIDAD": cantidad, "MATERIAL_COSTO_TOTAL": "",
                "FECHA_RESERVA": fecha_reserva_final, "EMPLAZAMIENTO": emplazamiento_final, "PROY_INSPECTOR_USUARIO": "", "TENSION_NIVEL": tension,
                "PROY_TIPOLOGIA_ID": params['tipologia'], "SECCION": seccion, "PROY_PARTE_ID": parte, 
                "CODIGO_INTERNO": codigo_row,
                "PROY_ZONA_ID": zona_final, 
                "PRESUPUESTO_FINALIDAD": fb_row,
                "ELEMENTO_PEP": "", "ODM_PRESUPUESTO": "", "ODM_PRESPTO_DESCRIPCION": descripcion, "TAM": ""
            }
            datos_finales_ot.append(fila_excel)

        df_final_ot = pd.DataFrame(datos_finales_ot)
        mostrar_tabla_bonita(df_final_ot)
        
        st.download_button(
            label="📥 Descargar Plantilla OT Completada (.xlsx)", data=generar_excel(df_final_ot),
            file_name=f"Plantilla_OT_{params['pry']}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

# --- SECCIÓN 4: Versionar Documentos ---
with tab4:
    st.subheader("📁 Versionar Documentos del Expediente")
    st.info("Sube tus documentos (Word, Excel, PDF). El sistema buscará el código antiguo y lo reemplazará por el nuevo, además de renombrar los archivos.")

    if 'ver_uploader_key' not in st.session_state: 
        st.session_state['ver_uploader_key'] = 0

    archivos_versionar = st.file_uploader("1. Sube los documentos (.docx, .xlsx, .pdf)", type=["docx", "xlsx", "pdf"], accept_multiple_files=True, key=f"file_ver_{st.session_state['ver_uploader_key']}")

    if not archivos_versionar:
        if 'zip_versionado' in st.session_state: del st.session_state['zip_versionado']
        if 'zip_name' in st.session_state: del st.session_state['zip_name']

    col_v1, col_v2 = st.columns(2)
    with col_v1:
        codigo_viejo = st.text_input("2. Código actual a buscar (ej: 212042026...)", key="cod_viejo")
    with col_v2:
        codigo_nuevo = st.text_input("3. Nuevo código de versión (ej: 800000021)", key="cod_nuevo")

    st.markdown("---")
    
    col_b1, col_b2 = st.columns([2, 8])
    btn_procesar_ver = col_b1.button("⚙️ Procesar Documentos y Versionar")

    # --- BOTÓN LIMPIAR VERSIONADOR ---
    if col_b2.button("🧹 Limpiar Todo", key="btn_limpiar_tab4"):
        st.session_state['ver_uploader_key'] += 1
        for k in ['cod_viejo', 'cod_nuevo', 'zip_versionado', 'zip_name']:
            st.session_state.pop(k, None)
        recargar_pagina()

    if btn_procesar_ver:
        if not archivos_versionar or not codigo_viejo or not codigo_nuevo:
            st.warning("⚠️ Sube al menos un archivo y completa ambos códigos.")
        else:
            with st.spinner("Procesando archivos... esto puede tomar un momento."):
                temp_dir = os.path.join(tempfile.gettempdir(), f"Procesado_{datetime.now().strftime('%H%M%S')}")
                os.makedirs(temp_dir, exist_ok=True)

                for archivo in archivos_versionar:
                    nombre_orig = archivo.name
                    nombre_sin_ext, ext = os.path.splitext(nombre_orig)
                    
                    nuevo_nombre = f"{codigo_nuevo}_{nombre_sin_ext}{ext}"
                    ruta_temp = os.path.join(temp_dir, nuevo_nombre)
                    ruta_absoluta = os.path.abspath(ruta_temp)

                    with open(ruta_absoluta, "wb") as f:
                        f.write(archivo.getbuffer())

                    os.chmod(ruta_absoluta, stat.S_IWRITE)

                    if ext.lower() == '.docx':
                        procesar_word(ruta_absoluta, codigo_viejo, codigo_nuevo)
                    elif ext.lower() == '.xlsx':
                        procesar_excel_automatico(ruta_absoluta, codigo_viejo, codigo_nuevo)
                    elif ext.lower() == '.pdf':
                        procesar_pdf(ruta_absoluta, codigo_viejo, codigo_nuevo)

                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
                    for root_dir, _, files in os.walk(temp_dir):
                        for file in files:
                            ruta_archivo = os.path.join(root_dir, file)
                            zipf.write(ruta_archivo, arcname=file)

                st.session_state['zip_versionado'] = zip_buffer.getvalue()
                st.session_state['zip_name'] = f"Expediente_{codigo_nuevo}.zip"

                try:
                    shutil.rmtree(temp_dir)
                except Exception:
                    pass

            st.success("✅ Archivos procesados con éxito.")

    if st.session_state.get('zip_versionado'):
        st.download_button(
            label="📥 Descargar Expediente Versionado (.ZIP)",
            data=st.session_state['zip_versionado'],
            file_name=st.session_state['zip_name'],
            mime="application/zip"
        )

# --- PORTADA FINAL AL FONDO ---
st.write("")
try:
    st.image(logo_path3, use_container_width=True)
except Exception:
    pass
